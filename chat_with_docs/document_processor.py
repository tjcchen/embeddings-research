import os
import requests
from typing import List, Dict, Any
from pathlib import Path

import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
from langchain_text_splitters import RecursiveCharacterTextSplitter, CharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument

from config import Config

class DocumentProcessor:
    def __init__(self):
        # Use simple CharacterTextSplitter to avoid recursion issues
        self.text_splitter = CharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separator="\n\n",
            length_function=len,
        )
        
        # Fallback simple splitter for problematic content
        self.simple_splitter = CharacterTextSplitter(
            chunk_size=Config.CHUNK_SIZE,
            chunk_overlap=Config.CHUNK_OVERLAP,
            separator="\n",
            length_function=len,
        )
        
    def process_file(self, file_path: str) -> List[LangChainDocument]:
        """Process a file and return chunked documents"""
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            text = self._extract_pdf_text(file_path)
        elif file_extension == '.docx':
            text = self._extract_docx_text(file_path)
        elif file_extension in ['.txt', '.md']:
            text = self._extract_text_file(file_path)
        elif file_extension == '.html':
            text = self._extract_html_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension}")
        
        # Create metadata
        metadata = {
            "source": file_path,
            "file_type": file_extension,
            "file_name": Path(file_path).name
        }
        
        # Split text into chunks with error handling
        documents = []
        try:
            chunks = self.text_splitter.split_text(text)
        except (RecursionError, Exception) as e:
            print(f"Primary text splitter failed, using fallback: {str(e)}")
            try:
                chunks = self.simple_splitter.split_text(text)
            except (RecursionError, Exception) as e2:
                print(f"Fallback splitter failed, using manual chunking: {str(e2)}")
                # Manual chunking as last resort
                chunk_size = Config.CHUNK_SIZE
                chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
        
        # Create LangChain documents
        for i, chunk in enumerate(chunks):
            if chunk.strip():  # Only add non-empty chunks
                doc_metadata = metadata.copy()
                doc_metadata["chunk_id"] = i
                documents.append(LangChainDocument(page_content=chunk.strip(), metadata=doc_metadata))
        
        return documents
    
    def process_url(self, url: str) -> List[LangChainDocument]:
        """Process a web page URL and return chunked documents"""
        try:
            # Add headers to avoid being blocked
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            
            response = requests.get(url, timeout=30, headers=headers)
            response.raise_for_status()
            
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract text
            text = soup.get_text()
            
            # Clean up text more carefully to avoid recursion issues
            lines = text.splitlines()
            cleaned_lines = []
            for line in lines:
                line = line.strip()
                if line and len(line) > 0:
                    cleaned_lines.append(line)
            
            text = '\n'.join(cleaned_lines)
            
            # Validate text length to prevent processing issues
            if len(text) > 1000000:  # 1MB limit
                text = text[:1000000]
                print(f"Warning: Text truncated to 1MB for URL: {url}")
            
            if not text.strip():
                raise Exception("No text content found in the webpage")
            
            # Create metadata
            metadata = {
                "source": url,
                "file_type": "web_page",
                "title": soup.title.string if soup.title else "Unknown"
            }
            
            # Split text into chunks with multiple fallback levels
            text_chunks = []
            try:
                text_chunks = self.text_splitter.split_text(text)
            except (RecursionError, Exception) as e:
                print(f"Primary text splitter failed for URL, using fallback: {str(e)}")
                try:
                    text_chunks = self.simple_splitter.split_text(text)
                except (RecursionError, Exception) as e2:
                    print(f"Fallback splitter failed for URL, using manual chunking: {str(e2)}")
                    # Manual chunking as last resort
                    chunk_size = Config.CHUNK_SIZE
                    text_chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]
            
            if not text_chunks:
                raise Exception("No text chunks generated from the webpage")
            
            # Create LangChain documents
            documents = []
            for i, chunk in enumerate(text_chunks):
                if chunk.strip():  # Only add non-empty chunks
                    doc_metadata = metadata.copy()
                    doc_metadata["chunk_id"] = i
                    documents.append(LangChainDocument(page_content=chunk.strip(), metadata=doc_metadata))
            
            if not documents:
                raise Exception("No valid documents created from the webpage")
            
            return documents
            
        except requests.RequestException as e:
            raise Exception(f"Network error processing URL {url}: {str(e)}")
        except RecursionError as e:
            raise Exception(f"Recursion error processing URL {url}. The webpage content may be too complex or contain circular references.")
        except Exception as e:
            raise Exception(f"Error processing URL {url}: {str(e)}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def _extract_html_text(self, file_path: str) -> str:
        """Extract text from HTML file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file.read(), 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            return soup.get_text()
    
    def get_supported_extensions(self) -> List[str]:
        """Return list of supported file extensions"""
        return Config.SUPPORTED_EXTENSIONS
